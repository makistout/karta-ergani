from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path
from datetime import date

OUT = Path(__file__).with_name("Προσφορά_erganiOS_Ενδεικτική_Εταιρεία.docx")
LOGO = Path(__file__).with_name("erganios-logo.png")

BLUE = "123A63"
CYAN = "18A7C6"
PALE = "EAF5F8"
LIGHT = "F2F4F7"
MID = "5E6B78"
WHITE = "FFFFFF"
INK = "182533"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, before, after, color in [("Heading 1",16,16,8,BLUE),("Heading 2",13,12,6,BLUE),("Heading 3",12,8,4,"1F4D78")]:
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=tcMar.find(qn(f"w:{m}"))
        if node is None: node=OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"),str(v)); node.set(qn("w:type"),"dxa")

def set_cell_width(cell, dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn("w:tcW"))
    if tcW is None: tcW=OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"),str(dxa)); tcW.set(qn("w:type"),"dxa")

def set_table_geometry(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn("w:tblW")); tblW.set(qn("w:w"),str(sum(widths))); tblW.set(qn("w:type"),"dxa")
    ind=OxmlElement("w:tblInd"); ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa"); tblPr.append(ind)
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc=OxmlElement("w:gridCol"); gc.set(qn("w:w"),str(w)); grid.append(gc)
    for row in table.rows:
        for cell,w in zip(row.cells,widths): set_cell_width(cell,w); cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_bullet(text, compact=False):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25)
    p.paragraph_format.space_after=Pt(2 if compact else 8); p.paragraph_format.line_spacing=1.0 if compact else 1.167
    r=p.add_run(text)
    if compact: r.font.size=Pt(9.5)

def add_callout(label, text):
    t=doc.add_table(rows=1, cols=1); set_table_geometry(t,[9360]); c=t.cell(0,0); shade(c,PALE)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label+"  "); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

# Header / footer
hp=sec.header.paragraphs[0]
if LOGO.exists(): hp.add_run().add_picture(str(LOGO), width=Inches(1.55))
hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=fp.add_run("erganiOS  |  info@erganios.gr  |  697 739 2742")
fr.font.size=Pt(9); fr.font.color.rgb=RGBColor.from_string(MID)

# Cover / customer pack
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(22); p.paragraph_format.space_after=Pt(4)
r=p.add_run("ΕΜΠΟΡΙΚΗ ΠΡΟΣΦΟΡΑ"); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(CYAN)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
r=p.add_run("erganiOS"); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(22)
r=p.add_run("Ψηφιακή Κάρτα Εργασίας, Απολογιστικό & AI Agent"); r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(MID)

meta=doc.add_table(rows=4, cols=2); meta.style="Table Grid"
rows=[("Προς","ΕΝΔΕΙΚΤΙΚΗ ΕΜΠΟΡΙΚΗ ΙΚΕ"),("Αντικείμενο","Παροχή υπηρεσίας erganiOS"),("Ημερομηνία","26 Αυγούστου 2026"),("Ισχύς προσφοράς","30 ημερολογιακές ημέρες")]
for i,(a,b) in enumerate(rows):
    meta.cell(i,0).text=a; meta.cell(i,1).text=b; shade(meta.cell(i,0),LIGHT)
    meta.cell(i,0).paragraphs[0].runs[0].bold=True
set_table_geometry(meta,[2200,7160])

doc.add_heading("Η πρότασή μας", level=1)
doc.add_paragraph("Η erganiOS προτείνει μία ενιαία cloud λύση για την καθημερινή διαχείριση της Ψηφιακής Κάρτας Εργασίας και των σχετικών αποκλίσεων. Η πρόταση οργανώνεται σε δύο εναλλακτικούς άξονες, ώστε η επιχείρηση να επιλέξει το επίπεδο αυτοματοποίησης που ανταποκρίνεται στις ανάγκες και στο μέγεθός της.")
add_callout("Κεντρικό όφελος", "Απολογιστικό με 3 κλικ, ομαδοποιημένα ζητήματα και σαφείς προτάσεις ενεργειών — χωρίς διαρκή εναλλαγή μεταξύ portal, αρχείων Excel και πολλαπλών οθονών.")

doc.add_heading("Τι καλύπτει το erganiOS", level=1)
for x in [
    "Ενιαία εικόνα εργαζομένων, δηλωμένων ωραρίων, πραγματικής απασχόλησης και εκκρεμοτήτων.",
    "Απολογιστικό και ωρομέτρηση με ομαδοποίηση αποκλίσεων και προτάσεις διόρθωσης.",
    "Ημερήσιες και εβδομαδιαίες αλλαγές ωραρίου, σταθερές εβδομάδες και οργάνωση βαρδιών.",
    "Εντοπισμός ελλιπών χτυπημάτων, καθυστερήσεων και εργασίας χωρίς δηλωμένο ωράριο.",
    "Άμεση εικόνα του ποιος εργάζεται τώρα και γρήγορες ενέργειες από την καθημερινή οθόνη.",
    "Ειδοποιήσεις μέσω Email ή Telegram προς τους κατάλληλους χρήστες.",
    "Cloud λειτουργία 24/7, εγκατάσταση σε λίγα λεπτά και χρήση χωρίς πρόσθετο εξοπλισμό.",
    "Συμβατότητα με το Π.Σ. ΕΡΓΑΝΗ ΙΙ και υποστήριξη πρακτικών σεναρίων Ψηφιακής Κάρτας."
]: add_bullet(x)

doc.add_page_break()
doc.add_heading("Άξονας 1 — erganiOS + Απολογιστικό", level=1)
doc.add_paragraph("Το βασικό πακέτο περιλαμβάνει το erganiOS, το Απολογιστικό και την Ωρομέτρηση. Η συνδρομή τιμολογείται ετησίως, ανάλογα με το πλήθος εργαζομένων.")

base=[("1–2","94 €"),("3–5","112 €"),("6–20","148 €"),("21–50","187 €"),("51–100","274 €"),("101–200","439 €")]
t=doc.add_table(rows=1, cols=3); t.style="Table Grid"
for j,h in enumerate(["Εργαζόμενοι","Διάρκεια","Ετήσια συνδρομή"]):
    t.cell(0,j).text=h; shade(t.cell(0,j),BLUE)
    for rr in t.cell(0,j).paragraphs[0].runs: rr.bold=True; rr.font.color.rgb=RGBColor.from_string(WHITE)
for rng,price in base:
    cells=t.add_row().cells; cells[0].text=rng; cells[1].text="1 έτος"; cells[2].text=price
    cells[1].paragraphs[0].alignment=cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
set_table_geometry(t,[3000,2200,4160])
p=doc.add_paragraph("Πηγή τιμών: δημοσιευμένος τιμοκατάλογος erganios.gr, έλεγχος 26/08/2026.")
p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4); p.runs[0].italic=True; p.runs[0].font.size=Pt(9); p.runs[0].font.color.rgb=RGBColor.from_string(MID)

doc.add_heading("Άξονας 2 — erganiOS + Απολογιστικό + AI Agent", level=1)
doc.add_paragraph("Στον δεύτερο άξονα προστίθεται AI Agent για αυτοματοποιημένες διαδικασίες της ημερήσιας λειτουργίας και της διαχείρισης καρτών. Η χρέωση του AI Agent είναι μηνιαία και προστίθεται στην αντίστοιχη ετήσια συνδρομή του βασικού πακέτου.")

ai=[("1–2","94 €","40 €","574 €"),("3–5","112 €","45 €","652 €"),("6–20","148 €","50 €","748 €"),("21–50","187 €","55 €","847 €"),("51–100","274 €","60 €","994 €"),("101–200","439 €","70 €","1.279 €")]
t=doc.add_table(rows=1, cols=4); t.style="Table Grid"
headers=["Εργαζόμενοι","erganiOS / έτος","AI Agent / μήνα","Σύνολο 12μήνου"]
for j,h in enumerate(headers):
    t.cell(0,j).text=h; shade(t.cell(0,j),CYAN)
    for rr in t.cell(0,j).paragraphs[0].runs: rr.bold=True; rr.font.color.rgb=RGBColor.from_string(WHITE)
for row in ai:
    cells=t.add_row().cells
    for j,val in enumerate(row): cells[j].text=val; cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
set_table_geometry(t,[2100,2380,2480,2400])
p=doc.add_paragraph()
p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(0)
r=p.add_run("Σημείωση υπολογισμού: "); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor.from_string(BLUE)
r=p.add_run("Το σύνολο 12μήνου ισούται με την ετήσια συνδρομή erganiOS συν 12 μηνιαίες χρεώσεις AI Agent. Η κλιμάκωση 40€–70€/μήνα αποτελεί την παρούσα εμπορική πρόταση."); r.font.size=Pt(9.5)

doc.add_heading("Ενδεικτικό σενάριο για την εταιρεία", level=1)
doc.add_paragraph("Για επιχείρηση με 21–50 εργαζομένους, οι δύο επιλογές αποτυπώνονται ως εξής:")
t=doc.add_table(rows=1, cols=3); t.style="Table Grid"
for j,h in enumerate(["Επιλογή","Χρέωση","Περιλαμβάνει"]):
    t.cell(0,j).text=h; shade(t.cell(0,j),BLUE)
    for rr in t.cell(0,j).paragraphs[0].runs: rr.bold=True; rr.font.color.rgb=RGBColor.from_string(WHITE)
data=[
    ("Άξονας 1","187 € / έτος","erganiOS, Απολογιστικό και Ωρομέτρηση"),
    ("Άξονας 2","187 € / έτος + 55 € / μήνα\n847 € / 12μηνο","Όλα τα παραπάνω και AI Agent")]
for a,b,c in data:
    cells=t.add_row().cells; cells[0].text=a; cells[1].text=b; cells[2].text=c
set_table_geometry(t,[2000,2860,4500])

doc.add_heading("Προτεινόμενη επιλογή", level=1)
doc.add_paragraph("Για επιχείρηση που επιδιώκει ενεργή αυτοματοποίηση, προτείνεται ο Άξονας 2. Ο AI Agent περιορίζει τη χειροκίνητη παρακολούθηση και επιταχύνει τη διαχείριση εκκρεμοτήτων της Ψηφιακής Κάρτας.")

doc.add_heading("Εμπορικοί όροι", level=1)
terms=[
    "Οι αναγραφόμενες τιμές δεν περιλαμβάνουν Φ.Π.Α. 24%.",
    "Η βασική συνδρομή erganiOS είναι ετήσια και προκαταβάλλεται με την ενεργοποίηση.",
    "Η πρόσθετη υπηρεσία AI Agent χρεώνεται μηνιαίως. Το 12μηνο σύνολο παρέχεται για εύκολη σύγκριση.",
    "Η τελική βαθμίδα καθορίζεται από το πλήθος ενεργών εργαζομένων κατά την ενεργοποίηση και επανεξετάζεται σε ουσιώδη μεταβολή του προσωπικού.",
    "Η έναρξη προϋποθέτει παροχή των απαραίτητων στοιχείων πρόσβασης και εξουσιοδοτήσεων από τον πελάτη.",
    "Τυχόν ειδικές διασυνδέσεις, παραμετροποιήσεις ή υπηρεσίες πέραν της τυπικής ενεργοποίησης συμφωνούνται χωριστά.",
    "Η παρούσα προσφορά είναι ενδεικτική, απευθύνεται σε υποθετική εταιρεία και ισχύει για 30 ημέρες από την έκδοσή της."
]
for x in terms: add_bullet(x, compact=True)

doc.add_heading("Αποδοχή προσφοράς", level=1)
doc.add_paragraph("Η αποδοχή τελεί υπό την προϋπόθεση υπογραφής της οριστικής σύμβασης παροχής υπηρεσιών και των όρων προστασίας προσωπικών δεδομένων.")

t=doc.add_table(rows=4, cols=2); t.style="Table Grid"
sig=[("Επιλεγμένος άξονας","☐ Άξονας 1    ☐ Άξονας 2"),("Επωνυμία πελάτη",""),("Ονοματεπώνυμο / Ιδιότητα", ""),("Ημερομηνία / Υπογραφή","__________ / ____________________")]
for i,(a,b) in enumerate(sig): t.cell(i,0).text=a; t.cell(i,1).text=b; shade(t.cell(i,0),LIGHT); t.cell(i,0).paragraphs[0].runs[0].bold=True
set_table_geometry(t,[3300,6060])

doc.core_properties.title="Εμπορική Προσφορά erganiOS"
doc.core_properties.subject="erganiOS + Απολογιστικό / erganiOS + AI Agent"
doc.core_properties.author="erganiOS"
doc.save(OUT)
print(OUT)
