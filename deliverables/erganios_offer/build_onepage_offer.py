from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

BASE = Path(__file__).parent
OUT = BASE / "Μονοσέλιδη_Προσφορά_erganiOS.docx"
LOGO = BASE / "erganios-logo.png"
BLUE, CYAN, PALE, LIGHT, MID, WHITE, INK = "123A63", "18A7C6", "EAF5F8", "F2F4F7", "5E6B78", "FFFFFF", "182533"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(.60)  # one-page commercial-offer override
sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.32)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(9.5); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(3); normal.paragraph_format.line_spacing = 1.0
for name, size, before, after in [("Heading 1",14,8,4),("Heading 2",11.5,6,3)]:
    s=doc.styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(BLUE)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); shd=pr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); pr.append(shd)
    shd.set(qn("w:fill"), fill)

def geom(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    pr=table._tbl.tblPr; tw=pr.find(qn("w:tblW")); tw.set(qn("w:w"),str(sum(widths))); tw.set(qn("w:type"),"dxa")
    ind=OxmlElement("w:tblInd"); ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa"); pr.append(ind)
    grid=table._tbl.tblGrid
    for x in list(grid): grid.remove(x)
    for w in widths:
        g=OxmlElement("w:gridCol"); g.set(qn("w:w"),str(w)); grid.append(g)
    for row in table.rows:
        for c,w in zip(row.cells,widths):
            cp=c._tc.get_or_add_tcPr(); cw=cp.find(qn("w:tcW"));
            if cw is None: cw=OxmlElement("w:tcW"); cp.append(cw)
            cw.set(qn("w:w"),str(w)); cw.set(qn("w:type"),"dxa")
            mar=OxmlElement("w:tcMar")
            for side,val in (("top",55),("start",90),("bottom",55),("end",90)):
                e=OxmlElement(f"w:{side}"); e.set(qn("w:w"),str(val)); e.set(qn("w:type"),"dxa"); mar.append(e)
            cp.append(mar); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def bullet(text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.32); p.paragraph_format.first_line_indent=Inches(-.16)
    p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.0; p.add_run(text)

# Compact customer-pack header
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
if LOGO.exists(): p.add_run().add_picture(str(LOGO), width=Inches(1.35))
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
r=p.add_run("ΕΜΠΟΡΙΚΗ ΠΡΟΣΦΟΡΑ"); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(CYAN)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
r=p.add_run("erganiOS | Ψηφιακή Κάρτα Εργασίας"); r.bold=True; r.font.size=Pt(21); r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
r=p.add_run("Προς: ΕΝΔΕΙΚΤΙΚΗ ΕΜΠΟΡΙΚΗ ΙΚΕ   •   Ημερομηνία: 27/08/2026"); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(MID)

t=doc.add_table(rows=1, cols=1); c=t.cell(0,0); shade(c,PALE); geom(t,[9360])
p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
r=p.add_run("Απολογιστικό με 3 κλικ. "); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
p.add_run("Ενιαία εικόνα ωραρίων και χτυπημάτων, ομαδοποιημένες αποκλίσεις και σαφείς προτάσεις ενεργειών, χωρίς διαρκή χρήση portal και Excel.")
p=c.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(0)
r=p.add_run("AI Agent για καθημερινή αυτοματοποίηση. "); r.bold=True; r.font.color.rgb=RGBColor.from_string(CYAN)
p.add_run("Παρακολουθεί τις εκκρεμότητες, υποστηρίζει τη διαχείριση καρτών και μειώνει τις επαναλαμβανόμενες χειροκίνητες ενέργειες.")

doc.add_heading("Τι περιλαμβάνει", level=1)
for text in [
    "Απολογιστικό και ωρομέτρηση με έλεγχο αποκλίσεων και προτάσεις διόρθωσης.",
    "Ημερήσιες/εβδομαδιαίες αλλαγές ωραρίου, σταθερές εβδομάδες και βάρδιες.",
    "Εντοπισμό ελλιπών χτυπημάτων, καθυστερήσεων και εργασίας χωρίς ωράριο.",
    "Live εικόνα απασχόλησης και ειδοποιήσεις μέσω Email ή Telegram.",
    "Cloud 24/7, γρήγορη εγκατάσταση, χωρίς πρόσθετο εξοπλισμό, συμβατό με ΕΡΓΑΝΗ ΙΙ."
]: bullet(text)

doc.add_heading("Δύο άξονες συνεργασίας", level=1)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
r=p.add_run("erganiOS + Απολογιστικό: "); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
p.add_run("ετήσια συνδρομή με Απολογιστικό και Ωρομέτρηση.  ")
r=p.add_run("erganiOS + Απολογιστικό + AI Agent: "); r.bold=True; r.font.color.rgb=RGBColor.from_string(CYAN)
p.add_run("το βασικό πακέτο μαζί με ετήσια υπηρεσία αυτοματοποίησης καθημερινών διαδικασιών και διαχείρισης καρτών.")

prices=[
    ("1–2","94 €","480 €","574 €"),("3–5","112 €","540 €","652 €"),("6–20","148 €","600 €","748 €"),
    ("21–50","187 €","660 €","847 €"),("51–100","274 €","720 €","994 €"),("101–200","439 €","840 €","1.279 €")]
t=doc.add_table(rows=1, cols=4); t.style="Table Grid"
for j,h in enumerate(["Εργαζόμενοι","erganiOS +\nΑπολογιστικό\nετησίως","AI Agent\nετησίως","erganiOS + Απολογιστικό\n+ AI Agent\nετησίως"]):
    t.cell(0,j).text=h; shade(t.cell(0,j),BLUE)
    for rr in t.cell(0,j).paragraphs[0].runs: rr.bold=True; rr.font.size=Pt(8.5); rr.font.color.rgb=RGBColor.from_string(WHITE)
    t.cell(0,j).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
for row in prices:
    cells=t.add_row().cells
    for j,val in enumerate(row):
        cells[j].text=val; cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        for rr in cells[j].paragraphs[0].runs: rr.font.size=Pt(8.5)
geom(t,[1900,2350,2350,2760])

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
r=p.add_run("Υπολογισμός πλήρους πακέτου: "); r.bold=True; r.font.size=Pt(8)
r=p.add_run("ετήσια συνδρομή erganiOS + Απολογιστικό και ετήσια χρέωση AI Agent. Όλες οι αναγραφόμενες χρεώσεις είναι ετήσιες."); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MID)

doc.add_heading("Εμπορικοί όροι", level=2)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
r=p.add_run("Οι τιμές δεν περιλαμβάνουν ΦΠΑ 24%. "); r.bold=True
p.add_run("Η συνδρομή και στα δύο πακέτα είναι ετήσια και προκαταβάλλεται. Η βαθμίδα καθορίζεται από τους ενεργούς εργαζομένους. Ειδικές διασυνδέσεις ή παραμετροποιήσεις συμφωνούνται χωριστά.")

doc.add_heading("Ειδική εμπορική προσφορά", level=2)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
r=p.add_run("Έκπτωση 25% "); r.bold=True; r.font.color.rgb=RGBColor.from_string(CYAN)
p.add_run("στις ετήσιες συνδρομές, στο πλαίσιο σύμβασης διάρκειας δύο (2) ετών.")
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
p.add_run("Κατά την έναρξη πραγματοποιούνται η αρχική παραμετροποίηση και η έκδοση προσωποποιημένων QR Codes εργαζομένων με το ονοματεπώνυμό τους. ")
r=p.add_run("Η αρχική εκπαίδευση στη χρήση του erganiOS παρέχεται δωρεάν."); r.bold=True
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8)
p.add_run("Τυχόν κόστος αρχικής προετοιμασίας και παραγωγής των QR Codes χρεώνεται εφάπαξ κατά την ενεργοποίηση της υπηρεσίας.")

t=doc.add_table(rows=2, cols=2); t.style="Table Grid"; geom(t,[4680,4680])
choices=[
    ("☐  Επιλέγω erganiOS + Απολογιστικό", "Υπογραφή: ____________________"),
    ("☐  Επιλέγω erganiOS + Απολογιστικό + AI Agent", "Υπογραφή: ____________________"),
]
for i,(choice,signature) in enumerate(choices):
    t.rows[i].height = Pt(34)
    t.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    p=t.cell(i,0).paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(choice); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(BLUE if i == 0 else CYAN)
    p=t.cell(i,1).paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(signature); r.font.size=Pt(8.5)

fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=fp.add_run("erganiOS  |  info@erganios.gr  |  697 739 2742  |  erganios.gr"); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MID)

doc.core_properties.title="Μονοσέλιδη Εμπορική Προσφορά erganiOS"
doc.core_properties.author="erganiOS"
doc.save(OUT)
print(OUT)
